"""文档加载与切块的极简实现。

仅支持 .txt / .md / .pdf / .docx 四种最常见格式。生产环境请见 Yuxi 的
backend/package/yuxi/plugins/parser/ —— 那里有 MinerU/PaddleX/OCR 多种解析器。
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


def _read_text(path: Path) -> str:
    """按文件扩展名分发读取。"""
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")

    if suffix == ".pdf":
        # PyMuPDF：纯 Python，比 pypdf 更稳，比 unstructured 更轻
        import pymupdf

        with pymupdf.open(path) as doc:
            return "\n\n".join(page.get_text() for page in doc)

    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise ValueError(f"不支持的文件类型：{suffix}（{path.name}）")


def load_docs_from_dir(dir_path: str | Path) -> list[Document]:
    """加载目录下所有支持的文件，每个文件一篇 Document。

    Document.metadata 至少包含 source（文件名）。
    """
    dir_path = Path(dir_path)
    if not dir_path.exists():
        raise FileNotFoundError(f"目录不存在：{dir_path}")

    documents: list[Document] = []
    for file_path in sorted(dir_path.iterdir()):
        if file_path.is_file() and file_path.suffix.lower() in {".txt", ".md", ".pdf", ".docx"}:
            content = _read_text(file_path)
            if content.strip():
                documents.append(
                    Document(
                        page_content=content,
                        metadata={"source": file_path.name},
                    )
                )
    return documents


def split_documents(
    documents: list[Document],
    chunk_size: int = 500,
    chunk_overlap: int = 80,
) -> list[Document]:
    """递归切块。中文友好的分隔符顺序：段落 → 行 → 句号 → 逗号 → 字。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    return splitter.split_documents(documents)
